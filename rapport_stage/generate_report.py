# -*- coding: utf-8 -*-
"""Génération du rapport de stage complet (Word .docx)."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
IMG = os.path.join(HERE, "img")
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
GREY = RGBColor(0x64, 0x74, 0x8B)

doc = Document()

# --------------------------------------------------------------- mise en page
sec = doc.sections[0]
sec.page_height = Cm(29.7); sec.page_width = Cm(21)
sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3); sec.right_margin = Cm(2.5)

st = doc.styles["Normal"]
st.font.name = "Times New Roman"; st.font.size = Pt(12)
st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for h, sz in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12.5)]:
    s = doc.styles[h]
    s.font.name = "Times New Roman"; s.font.size = Pt(sz)
    s.font.color.rgb = NAVY; s.font.bold = True


FIG_N = [0]; TAB_N = [0]
FIGURES = []; TABLES = []


def page_number_footer():
    footer = sec.footer
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    run._r.addprevious(fld)


def br():
    doc.add_page_break()


def h1(text):
    p = doc.add_heading(text, level=1); return p


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def para(text, bold=False, italic=False, align=None, size=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    if align == "c": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "r": p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet"); return p


def numbered(text):
    p = doc.add_paragraph(text, style="List Number"); return p


def figure(name, caption):
    FIG_N[0] += 1
    path = os.path.join(IMG, name)
    doc.add_picture(path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {FIG_N[0]} : {caption}")
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GREY
    FIGURES.append((FIG_N[0], caption))


def table(headers, rows, caption):
    TAB_N[0] += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cap.add_run(f"Tableau {TAB_N[0]} : {caption}")
    rr.italic = True; rr.font.size = Pt(10); rr.font.color.rgb = GREY
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True; run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    TABLES.append((TAB_N[0], caption))
    doc.add_paragraph()


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for line in text.strip("\n").split("\n"):
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # fond gris clair
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F4F6FB")
    p._p.get_or_add_pPr().append(shd)


def toc_field(title):
    h1(title)
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Faites « Mettre à jour les champs » (clic droit) pour générer le sommaire."
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    for el in (fldBegin, instr, fldSep, t, fldEnd):
        run._r.append(el)


# ============================================================== PAGE DE GARDE
def cover():
    para("BURKINA FASO", bold=True, align="c", size=12)
    para("Unité – Progrès – Justice", italic=True, align="c", size=11)
    para("MINISTÈRE DE L'ENSEIGNEMENT SUPÉRIEUR, DE LA RECHERCHE ET DE L'INNOVATION", align="c", size=10)
    doc.add_paragraph()
    para("[NOM DE L'UNIVERSITÉ / INSTITUT SUPÉRIEUR]", bold=True, align="c", size=13, color=NAVY)
    para("[Nom de l'École / UFR]", align="c", size=11)
    para("Filière : [Informatique / Génie Logiciel]", align="c", size=11)
    para("Cycle Licence — Niveau 3", align="c", size=11)
    for _ in range(2): doc.add_paragraph()
    para("RAPPORT DE STAGE DE FIN DE CYCLE", bold=True, align="c", size=14, color=INDIGO)
    para("En vue de l'obtention de la Licence en Informatique", italic=True, align="c", size=11)
    doc.add_paragraph()
    # cadre titre
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONCEPTION ET IMPLÉMENTATION D'UNE APPLICATION\n"
                  "DE GESTION DES DOSSIERS ET DE PLANIFICATION DES ACTIVITÉS\n"
                  "POUR UN CABINET D'HUISSIER DE JUSTICE")
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    for edge in ("top", "bottom", "left", "right"):
        b = OxmlElement(f"w:{edge}"); b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "12")
        b.set(qn("w:space"), "6"); b.set(qn("w:color"), "1E3A5F")
        pPr.append(b)
    for _ in range(2): doc.add_paragraph()
    # auteurs / encadreurs
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Présenté et soutenu par :"
    t.cell(0, 1).text = "Sous la direction de :"
    t.cell(1, 0).text = "[Prénom NOM de l'étudiant]\nMatricule : [……]"
    t.cell(1, 1).text = "Maître de stage : [Me ……], Huissier de Justice\nEncadreur pédagogique : [M./Mme ……]"
    for row in t.rows:
        for c in row.cells:
            for pp in c.paragraphs:
                for rn in pp.runs: rn.font.size = Pt(11)
    for _ in range(2): doc.add_paragraph()
    para("Structure d'accueil : Cabinet d'Huissier de Justice Me SAWADOGO — Ouagadougou", align="c", size=11, bold=True)
    para("Période de stage : [date début] au [date fin] — Année académique 2025–2026", align="c", size=10, italic=True)
    br()


page_number_footer()
cover()

# ============================================================== DÉDICACE
h1("DÉDICACE")
doc.add_paragraph()
para("Je dédie ce modeste travail :", italic=True)
doc.add_paragraph()
para("À mon père et à ma mère, pour leur amour inconditionnel, leurs sacrifices "
     "constants et le soutien moral et matériel qu'ils n'ont jamais cessé de m'apporter "
     "tout au long de mon parcours. Que ce travail soit le témoignage de ma profonde gratitude.")
para("À mes frères et sœurs, pour leurs encouragements et leur présence à mes côtés.")
para("À tous mes enseignants, qui ont contribué à ma formation et m'ont transmis le goût "
     "du savoir et de la rigueur.")
para("À mes camarades de promotion, avec qui j'ai partagé des moments inoubliables "
     "d'entraide et de persévérance.")
para("À toutes les personnes qui, de près ou de loin, ont contribué à la réalisation de ce travail.")
br()

# ============================================================== REMERCIEMENTS
h1("REMERCIEMENTS")
para("La réalisation de ce rapport de stage n'aurait pas été possible sans le concours "
     "de plusieurs personnes à qui je tiens à exprimer ma sincère reconnaissance.")
para("Je remercie tout d'abord le Tout-Puissant pour la santé, le courage et la volonté "
     "qu'Il m'a accordés durant toute cette période.")
para("Mes remerciements les plus chaleureux s'adressent à mon maître de stage, "
     "[Me ……], Huissier de Justice, pour m'avoir accueilli au sein de son cabinet, "
     "pour sa disponibilité, ses conseils avisés et la confiance qu'il m'a témoignée "
     "en me confiant la conception de cette application.")
para("Je tiens également à exprimer ma gratitude à mon encadreur pédagogique, "
     "[M./Mme ……], pour son accompagnement, ses orientations méthodologiques et "
     "sa rigueur scientifique qui ont guidé l'ensemble de ce travail.")
para("Je remercie l'ensemble du corps professoral et administratif de "
     "[NOM DE L'UNIVERSITÉ / INSTITUT] pour la qualité de la formation dispensée.")
para("Mes remerciements vont aussi à tout le personnel du cabinet — clercs, secrétaires "
     "et assistants — pour leur accueil, leur patience et les nombreuses explications "
     "sur le métier d'huissier qui ont nourri l'analyse des besoins.")
para("Enfin, je remercie ma famille et mes amis pour leur soutien indéfectible.")
br()

# ============================================================== RÉSUMÉ
h1("RÉSUMÉ")
para("Le présent rapport rend compte des travaux menés durant notre stage de fin de cycle "
     "Licence, effectué au sein d'un cabinet d'huissier de justice. Le métier d'huissier "
     "repose sur la gestion d'un volume important de dossiers, d'actes juridiques, de parties, "
     "de clients, de paiements et de rendez-vous, dont le suivi manuel ou semi-informatisé "
     "engendre des lenteurs, des risques d'erreurs et des difficultés de traçabilité. "
     "Notre travail a consisté à concevoir et à implémenter une application web de gestion "
     "des dossiers et de planification des activités, destinée à centraliser l'information, "
     "à automatiser les traitements répétitifs et à sécuriser l'accès aux données.")
para("La démarche d'analyse et de conception s'est appuyée sur le formalisme UML complété "
     "par la méthode Merise pour la modélisation des données. La solution a été développée "
     "selon une architecture trois-tiers : un frontend en Angular utilisant la bibliothèque "
     "PrimeNG, une API REST développée avec le framework FastAPI (Python), et une base de "
     "données relationnelle PostgreSQL. L'application intègre un système d'authentification "
     "par jetons JWT, une gestion fine des droits par rôles (RBAC), un journal d'audit et "
     "un tableau de bord décisionnel.")
para("Mots-clés : huissier de justice, gestion de dossiers, application web, UML, Merise, "
     "FastAPI, Angular, PostgreSQL, RBAC, tableau de bord.", italic=True)
doc.add_paragraph()
para("ABSTRACT", bold=True, color=NAVY)
para("This report presents the work carried out during our end-of-cycle internship at a "
     "judicial bailiff's office. The bailiff profession relies on managing a large number "
     "of case files, legal acts, parties, clients, payments and appointments, whose manual "
     "follow-up leads to delays, errors and traceability issues. Our work consisted in "
     "designing and implementing a web application for case management and activity planning, "
     "intended to centralize information, automate repetitive tasks and secure data access. "
     "The analysis relied on UML and the Merise method. The solution follows a three-tier "
     "architecture: an Angular/PrimeNG frontend, a FastAPI (Python) REST API, and a "
     "PostgreSQL database, with JWT authentication, role-based access control, an audit log "
     "and a decision-support dashboard.")
para("Keywords: judicial bailiff, case management, web application, UML, Merise, FastAPI, "
     "Angular, PostgreSQL, RBAC, dashboard.", italic=True)
br()

# ============================================================== SIGLES
h1("LISTE DES SIGLES ET ABRÉVIATIONS")
table(["Sigle", "Signification"], [
    ["API", "Application Programming Interface (interface de programmation)"],
    ["CRUD", "Create, Read, Update, Delete (Créer, Lire, Modifier, Supprimer)"],
    ["CSS", "Cascading Style Sheets"],
    ["FCFA", "Franc de la Communauté Financière Africaine"],
    ["HTML", "HyperText Markup Language"],
    ["HTTP(S)", "HyperText Transfer Protocol (Secure)"],
    ["IDE", "Integrated Development Environment"],
    ["JSON", "JavaScript Object Notation"],
    ["JWT", "JSON Web Token"],
    ["MCD", "Modèle Conceptuel de Données"],
    ["MLD", "Modèle Logique de Données"],
    ["MPD", "Modèle Physique de Données"],
    ["ORM", "Object-Relational Mapping"],
    ["RBAC", "Role-Based Access Control (contrôle d'accès basé sur les rôles)"],
    ["REST", "Representational State Transfer"],
    ["SGBD", "Système de Gestion de Base de Données"],
    ["SPA", "Single Page Application"],
    ["SQL", "Structured Query Language"],
    ["UML", "Unified Modeling Language"],
    ["UI / UX", "User Interface / User Experience"],
], "Liste des sigles et abréviations")
br()

# La liste des figures et tableaux + sommaire seront placés ici (champs auto)
toc_field("LISTE DES FIGURES")
para("(La liste détaillée des figures se met à jour automatiquement dans Word via "
     "Références ▸ Insérer une table des illustrations, légende « Figure ».)", italic=True, size=10, color=GREY)
br()
toc_field("LISTE DES TABLEAUX")
para("(La liste détaillée des tableaux se met à jour automatiquement dans Word via "
     "Références ▸ Insérer une table des illustrations, légende « Tableau ».)", italic=True, size=10, color=GREY)
br()
toc_field("SOMMAIRE")
br()

# ============================================================== INTRODUCTION
h1("INTRODUCTION GÉNÉRALE")
para("À l'ère de la transformation numérique, l'informatisation des processus métiers "
     "constitue un levier majeur de performance pour toutes les organisations, y compris "
     "les professions juridiques. L'huissier de justice, officier ministériel et auxiliaire "
     "de justice, occupe une place centrale dans le fonctionnement de l'appareil judiciaire : "
     "il signifie les actes, procède aux constats, recouvre les créances et exécute les "
     "décisions de justice. L'exercice de ce métier s'accompagne de la production et de la "
     "manipulation d'un volume considérable de documents et de dossiers, soumis à des délais "
     "légaux stricts dont le non-respect peut entraîner la nullité des actes.")
para("Dans de nombreux cabinets, la gestion de ces dossiers demeure essentiellement manuelle "
     "ou repose sur des outils bureautiques disparates (tableurs, documents isolés, registres "
     "papier). Cette situation engendre des difficultés réelles : dispersion de l'information, "
     "redondance des saisies, lenteur dans la recherche d'un dossier, risque de perte de "
     "documents, absence de tableau de bord et difficulté à respecter les échéances. La "
     "planification des activités (rendez-vous, audiences, significations) souffre des mêmes "
     "limites, faute d'un agenda partagé et centralisé.")
para("C'est dans ce contexte que s'inscrit notre stage de fin de cycle, dont le thème est : "
     "« Conception et implémentation d'une application de gestion des dossiers et de "
     "planification des activités pour un cabinet d'huissier de justice ». L'objectif est de "
     "doter le cabinet d'un outil logiciel moderne, centralisé et sécurisé, capable de "
     "couvrir l'ensemble de ses processus métiers.")
para("Problématique. — La question centrale qui guide ce travail peut se formuler ainsi : "
     "comment concevoir et mettre en œuvre un système d'information capable de centraliser la "
     "gestion des dossiers, d'automatiser le suivi des actes et des paiements, de planifier "
     "efficacement les activités et de garantir la sécurité et la traçabilité des données au "
     "sein d'un cabinet d'huissier de justice ?", bold=False)
para("De cette problématique découlent plusieurs questions secondaires : quels sont les "
     "besoins réels des différents acteurs du cabinet ? Quelle architecture logicielle adopter "
     "pour assurer évolutivité et maintenabilité ? Comment garantir la confidentialité des "
     "données sensibles tout en facilitant le travail collaboratif ?")
para("Objectif général. — Concevoir et réaliser une application web de gestion des dossiers "
     "et de planification des activités adaptée aux besoins d'un cabinet d'huissier de justice.")
para("Objectifs spécifiques :", bold=True)
for o in ["analyser l'existant et recueillir les besoins fonctionnels et non fonctionnels ;",
          "modéliser le système à l'aide d'UML et de la méthode Merise ;",
          "concevoir une base de données relationnelle cohérente et normalisée ;",
          "développer les modules de gestion (dossiers, clients, parties, actes, paiements, agenda, documents, archives) ;",
          "mettre en place un mécanisme d'authentification et de gestion des droits par rôles ;",
          "concevoir un tableau de bord décisionnel et un journal d'audit ;",
          "tester, sécuriser et préparer le déploiement de l'application."]:
    bullet(o)
para("Méthodologie. — Notre démarche a combiné une phase d'observation et d'entretiens au "
     "sein du cabinet, une phase d'analyse et de conception orientée objet (UML) appuyée par "
     "Merise pour les données, puis une phase de réalisation itérative et de tests.")
para("Plan du rapport. — Ce document s'articule autour de trois grandes parties. La première "
     "présente le cadre général du stage : la structure d'accueil, l'étude de l'existant et la "
     "problématique. La deuxième est consacrée à l'analyse et à la conception du système "
     "(spécification des besoins, modélisation UML et Merise, architecture). La troisième "
     "décrit la réalisation : environnement technique et modules informatiques, présentation "
     "des interfaces, tests, sécurité et déploiement. Une conclusion générale dresse le bilan "
     "et ouvre des perspectives.")
br()

# ============================================================== PARTIE I
para("PREMIÈRE PARTIE", bold=True, align="c", size=16, color=INDIGO)
para("CADRE GÉNÉRAL DU STAGE", bold=True, align="c", size=14, color=NAVY)
doc.add_paragraph()
para("Cette première partie situe le contexte du stage. Elle présente le métier d'huissier "
     "de justice et la structure d'accueil, analyse de manière critique le système de gestion "
     "existant et dégage la problématique ainsi que les objectifs du projet.", italic=True, align="c")
br()

# ---- Chapitre 1
h1("CHAPITRE 1 : PRÉSENTATION DE LA STRUCTURE D'ACCUEIL")
h2("1.1. Le métier d'huissier de justice")
para("L'huissier de justice est un officier ministériel et public, titulaire d'une charge, "
     "qui exerce une profession libérale réglementée. Auxiliaire de justice, il joue un rôle "
     "essentiel dans le bon fonctionnement du système judiciaire. Ses missions principales "
     "peuvent être regroupées en plusieurs catégories.")
para("La signification des actes consiste à porter officiellement à la connaissance d'une "
     "personne un acte de procédure (assignation, jugement, commandement) en respectant des "
     "formes et des délais légaux précis. Le constat permet d'établir, de manière neutre et "
     "incontestable, la matérialité de faits (dégâts, occupation des lieux, contenu d'un site "
     "internet) ; le procès-verbal de constat fait foi jusqu'à preuve contraire. Le "
     "recouvrement de créances, amiable ou judiciaire, vise à obtenir le paiement des sommes "
     "dues à un créancier. Enfin, l'exécution des décisions de justice (saisies, expulsions) "
     "constitue le prolongement concret du jugement rendu.")
para("Chacune de ces missions se matérialise par l'ouverture d'un dossier, la rédaction "
     "d'actes, l'identification des parties (demandeur, défendeur, tiers), la perception "
     "d'émoluments et de provisions, et la planification de rendez-vous ou d'opérations sur "
     "le terrain. C'est précisément cette activité, fortement documentaire et soumise à des "
     "délais, qui justifie le besoin d'un système d'information dédié.")
h3("1.1.1. Cadre juridique de la profession")
para("L'huissier de justice exerce dans un cadre légal et déontologique strict. Sa "
     "nomination, ses compétences territoriales, sa responsabilité et la tarification de ses "
     "actes sont encadrées par des textes réglementaires. En sa qualité d'officier public, il "
     "confère l'authenticité aux actes qu'il dresse : ceux-ci font foi jusqu'à inscription de "
     "faux pour les énonciations qu'il a personnellement constatées. Cette force probante "
     "particulière impose une rigueur absolue dans la rédaction, la conservation et la "
     "traçabilité des actes — exigences que le système d'information doit impérativement "
     "soutenir.")
para("Par ailleurs, l'huissier est tenu au secret professionnel et à une obligation de "
     "confidentialité sur les affaires dont il a connaissance. Cette contrainte se traduit "
     "directement, dans notre application, par la nécessité d'un contrôle d'accès rigoureux et "
     "d'une journalisation des consultations et modifications.")
h3("1.1.2. Notion de délais et d'échéances")
para("La procédure civile repose largement sur des délais : délai pour signifier un jugement, "
     "délai d'appel, délai de paiement accordé à un débiteur, etc. Le non-respect de ces délais "
     "peut entraîner la forclusion ou la nullité. Le suivi rigoureux des échéances constitue "
     "donc un enjeu central, ce qui motive l'intégration, à terme, d'un système d'alertes dans "
     "l'application.")
h2("1.2. Historique et identité du cabinet")
para("Le cabinet d'huissier de justice Me SAWADOGO, situé à Ouagadougou, constitue notre "
     "structure d'accueil. Étude de taille moyenne, le cabinet emploie plusieurs collaborateurs "
     "(clercs, secrétaires, assistants) et traite quotidiennement des dossiers variés relevant "
     "tant du recouvrement que de la signification et du constat. Le cabinet est attaché à la "
     "qualité du service rendu à sa clientèle, composée de particuliers, d'entreprises, "
     "d'établissements bancaires et d'institutions.")
para("[Les informations relatives à la date de création, à la localisation exacte et à "
     "l'effectif sont à compléter par l'étudiant selon les données réelles du cabinet.]",
     italic=True, color=GREY, size=10)
h2("1.3. Missions et organisation")
para("Sur le plan organisationnel, le cabinet est dirigé par l'huissier titulaire, qui "
     "engage sa responsabilité sur l'ensemble des actes. Il s'appuie sur une équipe dont les "
     "rôles sont complémentaires : les clercs préparent et rédigent les actes, assurent les "
     "significations et le suivi des procédures ; les secrétaires gèrent l'accueil, le courrier "
     "et la prise de rendez-vous ; les assistants et agents de saisie assurent le suivi "
     "administratif, la comptabilité des provisions et l'archivage. L'organigramme suivant "
     "illustre cette structure.")
figure("organigramme.png", "Organigramme type d'un cabinet d'huissier de justice")
para("Cette organisation met en évidence la diversité des profils d'utilisateurs qui "
     "interagiront avec l'application. Elle justifie la mise en place d'une gestion des droits "
     "par rôles, afin que chaque collaborateur n'accède qu'aux fonctionnalités correspondant à "
     "ses attributions.")
h2("1.4. Déroulement et planning du stage")
para("Le stage s'est déroulé en plusieurs phases successives, depuis la prise de contact "
     "et la découverte du métier jusqu'à la rédaction du présent rapport, en passant par "
     "l'analyse des besoins, la conception et le développement. Le diagramme de Gantt "
     "ci-dessous synthétise l'ordonnancement et la durée des principales tâches.")
figure("gantt.png", "Diagramme de Gantt — planning de réalisation du stage")
para("Conclusion partielle. — Ce premier chapitre a permis de présenter le contexte "
     "professionnel du stage et de comprendre la nature de l'activité du cabinet. Le chapitre "
     "suivant analyse de manière critique le système de gestion en place.")
br()

# ---- Chapitre 2
h1("CHAPITRE 2 : ÉTUDE DE L'EXISTANT ET PROBLÉMATIQUE")
h2("2.1. Analyse de l'existant")
para("Avant toute conception, il est indispensable d'étudier le fonctionnement actuel du "
     "cabinet afin d'en comprendre les forces et les faiblesses. Cette étude a été menée par "
     "observation directe du travail quotidien et par des entretiens avec les différents "
     "collaborateurs.")
para("Il ressort de cette analyse que la gestion des dossiers s'effectue principalement à "
     "l'aide de registres papier et de fichiers bureautiques (tableurs et documents texte) "
     "stockés localement. Chaque dossier physique regroupe les pièces relatives à une affaire ; "
     "un numéro lui est attribué manuellement. Les rendez-vous sont notés sur un agenda papier "
     "ou dans des agendas électroniques individuels non partagés. Le suivi des paiements et des "
     "provisions est consigné dans un cahier ou un tableur comptable.")
h2("2.2. Critique de l'existant")
para("Cette organisation, bien que fonctionnelle, présente des limites importantes qui "
     "nuisent à l'efficacité et à la fiabilité du cabinet :")
for c in ["dispersion et redondance de l'information, une même donnée (coordonnées d'un client, "
          "par exemple) étant ressaisie dans plusieurs supports ;",
          "lenteur de la recherche d'un dossier ou d'un acte, particulièrement lorsque le volume "
          "d'archives augmente ;",
          "risque de perte, de détérioration ou de confusion entre documents papier ;",
          "absence de vision consolidée de l'activité (nombre de dossiers en cours, encaissements "
          "du mois, rendez-vous de la semaine) ;",
          "difficulté à respecter les délais légaux faute d'un système d'alerte sur les échéances ;",
          "absence de traçabilité : il est impossible de savoir qui a créé ou modifié une "
          "information et quand ;",
          "sécurité insuffisante des données confidentielles et difficulté à gérer les droits "
          "d'accès du personnel."]:
    bullet(c)
h2("2.3. Problématique")
para("Au regard de ces constats, la problématique se pose en ces termes : comment doter le "
     "cabinet d'un système d'information unique, fiable et sécurisé, permettant de centraliser "
     "la gestion des dossiers, d'automatiser le suivi des actes, des parties et des paiements, "
     "de planifier les activités et de garantir la traçabilité ainsi que la confidentialité des "
     "données ?")
h2("2.4. Objectifs du projet")
para("Pour répondre à cette problématique, le projet poursuit l'objectif général de concevoir "
     "et de réaliser une application web complète. Les objectifs spécifiques se déclinent en "
     "modules fonctionnels : gestion des dossiers, des clients, des parties, des actes, des "
     "paiements, de l'agenda, des documents et des archives, le tout encadré par une gestion "
     "des utilisateurs et des rôles, un journal d'audit et un tableau de bord.")
h2("2.5. Résultats attendus")
para("À l'issue du projet, le cabinet doit disposer d'une application accessible via un "
     "navigateur, offrant une saisie unique et centralisée des données, une recherche rapide, "
     "des statistiques en temps réel, une planification partagée des activités et un contrôle "
     "rigoureux des accès. La solution doit être évolutive, maintenable et conforme aux bonnes "
     "pratiques de sécurité.")
para("Conclusion partielle. — L'étude de l'existant a confirmé la pertinence du projet et "
     "précisé ses objectifs. La deuxième partie aborde l'analyse détaillée des besoins et la "
     "conception du système.")
br()

# ============================================================== PARTIE II
para("DEUXIÈME PARTIE", bold=True, align="c", size=16, color=INDIGO)
para("ANALYSE ET CONCEPTION DU SYSTÈME", bold=True, align="c", size=14, color=NAVY)
doc.add_paragraph()
para("Cette partie constitue le cœur méthodologique du travail. Elle spécifie les besoins, "
     "présente la modélisation du système à l'aide d'UML et de Merise, puis décrit "
     "l'architecture retenue pour la solution.", italic=True, align="c")
br()

# ---- Chapitre 3
h1("CHAPITRE 3 : SPÉCIFICATION DES BESOINS")
h2("3.1. Besoins fonctionnels")
para("Les besoins fonctionnels décrivent les services que le système doit rendre. Ils ont "
     "été regroupés par module.")
table(["Module", "Principales fonctionnalités"], [
    ["Authentification", "Connexion sécurisée, déconnexion, gestion du jeton de session"],
    ["Utilisateurs & rôles", "Créer/modifier/supprimer des comptes, attribuer des rôles et permissions"],
    ["Dossiers", "Ouvrir, consulter, modifier, classer et clôturer un dossier ; recherche et filtres"],
    ["Clients", "Gérer la fiche des clients (particuliers et personnes morales)"],
    ["Parties", "Enregistrer les parties d'un dossier (demandeur, défendeur, tiers)"],
    ["Actes", "Rédiger et suivre les actes juridiques, enregistrer leur résultat"],
    ["Paiements", "Enregistrer les paiements, provisions et émoluments ; journal des encaissements"],
    ["Agenda", "Planifier rendez-vous, audiences et significations ; vue hebdomadaire"],
    ["Documents", "Joindre, classer et télécharger les pièces d'un dossier"],
    ["Archives", "Archiver les dossiers clôturés et les retrouver"],
    ["Audit", "Journaliser toutes les actions sensibles des utilisateurs"],
    ["Tableau de bord", "Visualiser des indicateurs et statistiques en temps réel"],
], "Récapitulatif des besoins fonctionnels par module")
para("Au-delà de cette synthèse, chaque module porte des exigences précises. Le module "
     "d'authentification doit vérifier les identifiants, ouvrir une session sécurisée et la "
     "fermer à la déconnexion. Le module de gestion des dossiers doit permettre d'attribuer un "
     "numéro unique, de classer un dossier par type et par statut, de le rechercher et de "
     "suivre son évolution jusqu'à la clôture puis l'archivage. Le module de gestion des "
     "paiements doit autoriser l'enregistrement des règlements, calculer des totaux et restituer "
     "un journal des encaissements. Le module agenda doit offrir une planification partagée et "
     "une vue hebdomadaire. Le module d'audit doit journaliser, sans exception, toute action "
     "modifiant les données. Le tableau de bord doit présenter, en temps réel et selon le rôle "
     "de l'utilisateur, les indicateurs pertinents.")
para("Ces exigences fonctionnelles ont été hiérarchisées par priorité (indispensable, "
     "souhaitable, optionnel) afin de guider le développement itératif et de garantir que le "
     "cœur métier soit livré en premier.")
h2("3.2. Besoins non fonctionnels")
para("Au-delà des fonctionnalités, le système doit satisfaire un ensemble d'exigences de "
     "qualité :")
for b in ["Sécurité : authentification robuste, chiffrement des mots de passe, contrôle d'accès "
          "par rôles et journalisation des actions ;",
          "Performance : temps de réponse rapide pour la consultation et la recherche ;",
          "Convivialité (UX) : interface claire, cohérente et intuitive, adaptée aux différents profils ;",
          "Fiabilité et intégrité : cohérence des données garantie par le SGBD et les validations ;",
          "Maintenabilité : code structuré, modulaire et documenté ;",
          "Portabilité : application accessible depuis un navigateur, indépendamment du poste ;",
          "Évolutivité : architecture permettant l'ajout futur de nouveaux modules."]:
    bullet(b)
h2("3.3. Identification des acteurs")
para("Un acteur représente un rôle joué par un utilisateur (ou un système externe) qui "
     "interagit avec l'application. Quatre profils principaux ont été identifiés, conformément "
     "à l'organisation du cabinet et aux rôles implémentés dans la solution.")
table(["Acteur", "Description et droits"], [
    ["Administrateur", "Gère les utilisateurs, les rôles et la configuration ; accès complet, y compris à l'audit"],
    ["Huissier", "Supervise les dossiers, les actes, les paiements et le tableau de bord ; accès étendu"],
    ["Clerc", "Rédige les actes, gère les dossiers, les parties et les clients"],
    ["Secrétaire / Assistant", "Gère l'accueil, les clients, l'agenda et les documents ; accès restreint"],
], "Acteurs du système et niveaux d'accès")
h2("3.4. Diagramme de cas d'utilisation")
para("Le diagramme de cas d'utilisation offre une vue globale des fonctionnalités du système "
     "et des interactions avec les acteurs. Il met en évidence les services accessibles à "
     "chaque profil.")
figure("usecase.png", "Diagramme de cas d'utilisation global")
h3("3.4.1. Description textuelle de quelques cas d'utilisation")
para("Pour préciser le comportement attendu, certains cas d'utilisation sont décrits sous "
     "forme de fiches détaillant les scénarios nominaux et alternatifs.")
table(["Cas d'utilisation", "S'authentifier"], [
    ["Acteur principal", "Tout utilisateur enregistré"],
    ["Précondition", "L'utilisateur possède un compte actif"],
    ["Scénario nominal", "1. L'utilisateur saisit son email et son mot de passe ; 2. le système "
     "vérifie les identifiants ; 3. un jeton JWT est généré ; 4. l'utilisateur est redirigé vers "
     "le tableau de bord."],
    ["Scénario alternatif", "Identifiants invalides : un message d'erreur est affiché et l'accès refusé"],
    ["Postcondition", "Une session sécurisée est ouverte"],
], "Fiche descriptive du cas d'utilisation « S'authentifier »")
table(["Cas d'utilisation", "Créer un dossier"], [
    ["Acteur principal", "Clerc, Huissier"],
    ["Précondition", "L'utilisateur est authentifié et dispose des droits requis"],
    ["Scénario nominal", "1. L'utilisateur ouvre le formulaire de création ; 2. il saisit l'objet, "
     "le type, le client et les parties ; 3. le système valide et enregistre le dossier ; 4. un "
     "numéro unique est attribué et l'action est journalisée."],
    ["Scénario alternatif", "Données incomplètes : le système signale les champs obligatoires manquants"],
    ["Postcondition", "Le dossier est créé et visible dans la liste"],
], "Fiche descriptive du cas d'utilisation « Créer un dossier »")
table(["Cas d'utilisation", "Enregistrer un paiement"], [
    ["Acteur principal", "Huissier, Administrateur"],
    ["Précondition", "Le dossier concerné existe ; l'utilisateur dispose des droits financiers"],
    ["Scénario nominal", "1. L'utilisateur sélectionne le dossier ; 2. il saisit le montant, le "
     "type et le mode de paiement ; 3. le système valide et enregistre le paiement ; 4. les "
     "totaux et le tableau de bord sont actualisés."],
    ["Scénario alternatif", "Montant invalide : le système refuse l'enregistrement et le signale"],
    ["Postcondition", "Le paiement est comptabilisé et tracé"],
], "Fiche descriptive du cas d'utilisation « Enregistrer un paiement »")
table(["Cas d'utilisation", "Planifier un rendez-vous"], [
    ["Acteur principal", "Secrétaire, Clerc, Huissier"],
    ["Précondition", "L'utilisateur est authentifié"],
    ["Scénario nominal", "1. L'utilisateur ouvre l'agenda ; 2. il crée un événement avec titre, "
     "type, dates et lieu ; 3. il le rattache éventuellement à un dossier ; 4. le système "
     "enregistre et affiche l'événement dans la vue hebdomadaire."],
    ["Scénario alternatif", "Conflit d'horaire : le système avertit l'utilisateur"],
    ["Postcondition", "Le rendez-vous figure dans l'agenda partagé"],
], "Fiche descriptive du cas d'utilisation « Planifier un rendez-vous »")
para("Conclusion partielle. — La spécification des besoins ayant été établie, le chapitre "
     "suivant présente la modélisation détaillée du système.")
br()

# ---- Chapitre 4
h1("CHAPITRE 4 : MÉTHODOLOGIE ET MODÉLISATION")
h2("4.1. Choix de la méthode de conception")
para("La modélisation d'un système d'information peut s'appuyer sur différentes approches. "
     "Nous avons opté pour une démarche hybride associant UML (Unified Modeling Language) pour "
     "la modélisation orientée objet du comportement et de la structure du logiciel, et Merise "
     "pour la modélisation des données. UML, langage de modélisation standard, offre une "
     "panoplie de diagrammes (cas d'utilisation, séquence, activité, classes) particulièrement "
     "adaptés à une conception objet et à un développement moderne. Merise, méthode d'analyse "
     "française largement enseignée, apporte une rigueur reconnue dans la conception des bases "
     "de données à travers ses modèles conceptuel, logique et physique.")
para("Cette combinaison permet de tirer parti des forces des deux approches : la clarté "
     "fonctionnelle d'UML et la robustesse de Merise pour la structuration des données.")
h2("4.2. Diagramme de séquence")
para("Le diagramme de séquence décrit les interactions entre les objets du système au cours "
     "du temps, dans le cadre d'un scénario précis. Nous présentons ici deux scénarios "
     "représentatifs : l'authentification et la création d'un dossier.")
figure("sequence_auth.png", "Diagramme de séquence — authentification d'un utilisateur")
para("Lors de l'authentification, l'interface Angular transmet les identifiants à l'API "
     "FastAPI, qui les vérifie auprès du service d'authentification et de la base de données. "
     "En cas de succès, un jeton JWT signé est renvoyé puis stocké côté client pour sécuriser "
     "les requêtes ultérieures.")
figure("sequence_dossier.png", "Diagramme de séquence — création d'un dossier")
para("La création d'un dossier illustre le flux complet d'une opération d'écriture : "
     "validation des données par les schémas Pydantic, insertion en base via l'ORM, "
     "journalisation dans l'audit et retour d'une confirmation à l'utilisateur.")
h2("4.3. Diagramme d'activité")
para("Le diagramme d'activité modélise l'enchaînement des actions et des décisions. Le "
     "diagramme ci-dessous représente le cycle de vie d'un dossier, depuis la connexion de "
     "l'utilisateur jusqu'à l'archivage du dossier clôturé.")
figure("activite.png", "Diagramme d'activité — cycle de vie d'un dossier")
h2("4.4. Diagramme de classes")
para("Le diagramme de classes constitue la pièce maîtresse de la modélisation structurelle. "
     "Il décrit les classes du domaine, leurs attributs et les associations qui les relient. "
     "Il sert de fondement direct à la conception de la base de données.")
figure("classes.png", "Diagramme de classes du domaine")
para("Les classes principales sont : Cabinet, Utilisateur, Dossier, Client, Partie, Acte, "
     "Paiement, Agenda, Document et AuditLog. Le Dossier occupe une position centrale : il est "
     "rattaché à un cabinet, géré par un utilisateur, associé à un client, et regroupe des "
     "parties, des actes, des paiements, des rendez-vous et des documents.")
h2("4.5. Modèle Conceptuel de Données (MCD)")
para("Conformément à la méthode Merise, le Modèle Conceptuel de Données représente, de "
     "manière indépendante de toute technologie, les entités du système, leurs propriétés et "
     "les associations qui les unissent, assorties de cardinalités.")
figure("mcd.png", "Modèle Conceptuel de Données (formalisme Merise)")
para("Les cardinalités traduisent les règles de gestion : un dossier appartient à un et un "
     "seul cabinet (1,1) tandis qu'un cabinet peut contenir plusieurs dossiers (1,n) ; un "
     "dossier concerne un client et peut impliquer plusieurs parties ; un dossier peut donner "
     "lieu à plusieurs paiements et contenir plusieurs actes.")
h2("4.6. Modèle Logique de Données (MLD)")
para("Le passage du MCD au MLD applique les règles de transformation de Merise : chaque "
     "entité devient une table, chaque association de type un-à-plusieurs se traduit par une "
     "clé étrangère, et les identifiants deviennent des clés primaires. Le MLD relationnel "
     "obtenu est le suivant (les clés primaires sont soulignées et les clés étrangères "
     "préfixées par #).")
for line in [
    "CABINET (id_cabinet, nom, adresse, telephone, email)",
    "UTILISATEUR (id_user, nom, prenom, email, mot_de_passe, role, #id_cabinet)",
    "CLIENT (id_client, nom, type_client, telephone, email, adresse, #id_cabinet)",
    "DOSSIER (id_dossier, numero_dossier, objet, type_dossier, statut, date_creation, "
    "#id_cabinet, #id_user, #id_client)",
    "PARTIE (id_partie, nom, role_partie, adresse, #id_dossier)",
    "ACTE (id_acte, type_acte, date_acte, resultat, observations, #id_dossier)",
    "PAIEMENT (id_paiement, montant, type_paiement, mode_paiement, date_paiement, #id_dossier)",
    "AGENDA (id_agenda, titre, type_rdv, date_debut, date_fin, lieu, #id_dossier, #id_user)",
    "DOCUMENT (id_document, nom_fichier, type, chemin, date_ajout, #id_dossier)",
    "ARCHIVE (id_archive, raison_archivage, date_archivage, #id_dossier)",
    "AUDIT_LOG (id, action, entity_type, entity_id, description, date_action, #id_user)",
]:
    p = doc.add_paragraph(line)
    p.paragraph_format.left_indent = Cm(0.5)
    for r in p.runs: r.font.name = "Consolas"; r.font.size = Pt(10)
h2("4.7. Modèle Physique et dictionnaire de données")
para("Le Modèle Physique de Données précise, pour chaque champ, son type, sa taille et ses "
     "contraintes en vue de l'implémentation dans le SGBD PostgreSQL. Le dictionnaire de "
     "données ci-dessous en présente un extrait pour la table DOSSIER, table centrale du "
     "système.")
table(["Champ", "Type", "Contrainte", "Description"], [
    ["id_dossier", "SERIAL", "Clé primaire", "Identifiant unique du dossier"],
    ["numero_dossier", "VARCHAR(50)", "Unique, non nul", "Référence du dossier (ex. DOS-2026-0001)"],
    ["objet", "VARCHAR(255)", "Non nul", "Objet de l'affaire"],
    ["type_dossier", "VARCHAR / ENUM", "Non nul", "Recouvrement, constat, signification…"],
    ["statut", "VARCHAR / ENUM", "Non nul", "Nouveau, en cours, terminé, archivé, annulé"],
    ["date_creation", "TIMESTAMP", "Non nul", "Date d'ouverture du dossier"],
    ["id_client", "INTEGER", "Clé étrangère", "Référence du client concerné"],
    ["id_user", "INTEGER", "Clé étrangère", "Utilisateur ayant créé le dossier"],
], "Extrait du dictionnaire de données — table DOSSIER")
para("Le dictionnaire se poursuit pour les autres tables majeures du système. Les tableaux "
     "suivants présentent les structures des tables UTILISATEUR, CLIENT, PAIEMENT, ACTE et "
     "AGENDA.")
table(["Champ", "Type", "Contrainte", "Description"], [
    ["id_user", "SERIAL", "Clé primaire", "Identifiant unique de l'utilisateur"],
    ["nom / prenom", "VARCHAR(100)", "Non nul", "Identité du collaborateur"],
    ["email", "VARCHAR(150)", "Unique, non nul", "Identifiant de connexion"],
    ["mot_de_passe", "VARCHAR(255)", "Non nul", "Empreinte bcrypt du mot de passe"],
    ["role", "VARCHAR / ENUM", "Non nul", "ADMIN, HUISSIER, CLERC, SECRETAIRE, ASSISTANT"],
    ["id_cabinet", "INTEGER", "Clé étrangère", "Cabinet de rattachement"],
], "Dictionnaire de données — table UTILISATEUR")
table(["Champ", "Type", "Contrainte", "Description"], [
    ["id_client", "SERIAL", "Clé primaire", "Identifiant unique du client"],
    ["nom", "VARCHAR(150)", "Non nul", "Nom ou raison sociale"],
    ["type_client", "VARCHAR / ENUM", "Non nul", "Personne physique ou morale"],
    ["telephone", "VARCHAR(30)", "—", "Numéro de téléphone"],
    ["email", "VARCHAR(150)", "—", "Adresse électronique"],
    ["adresse", "VARCHAR(255)", "—", "Adresse postale"],
], "Dictionnaire de données — table CLIENT")
table(["Champ", "Type", "Contrainte", "Description"], [
    ["id_paiement", "SERIAL", "Clé primaire", "Identifiant unique du paiement"],
    ["montant", "NUMERIC(12,2)", "Non nul", "Montant perçu (FCFA)"],
    ["type_paiement", "VARCHAR / ENUM", "Non nul", "Provision, émolument, frais, solde"],
    ["mode_paiement", "VARCHAR / ENUM", "Non nul", "Espèces, virement, chèque, mobile money"],
    ["date_paiement", "DATE", "Non nul", "Date du règlement"],
    ["id_dossier", "INTEGER", "Clé étrangère", "Dossier concerné"],
], "Dictionnaire de données — table PAIEMENT")
table(["Champ", "Type", "Contrainte", "Description"], [
    ["id_acte", "SERIAL", "Clé primaire", "Identifiant unique de l'acte"],
    ["type_acte", "VARCHAR / ENUM", "Non nul", "Nature de l'acte"],
    ["date_acte", "DATE", "Non nul", "Date de l'acte"],
    ["resultat", "TEXT", "—", "Résultat ou issue de l'acte"],
    ["observations", "TEXT", "—", "Observations complémentaires"],
    ["id_dossier", "INTEGER", "Clé étrangère", "Dossier de rattachement"],
], "Dictionnaire de données — table ACTE")
table(["Champ", "Type", "Contrainte", "Description"], [
    ["id_agenda", "SERIAL", "Clé primaire", "Identifiant unique du rendez-vous"],
    ["titre", "VARCHAR(150)", "Non nul", "Intitulé de l'activité"],
    ["type_rdv", "VARCHAR / ENUM", "Non nul", "Rendez-vous, audience, signification…"],
    ["date_debut", "TIMESTAMP", "Non nul", "Date et heure de début"],
    ["date_fin", "TIMESTAMP", "Non nul", "Date et heure de fin"],
    ["lieu", "VARCHAR(255)", "—", "Lieu de l'activité"],
], "Dictionnaire de données — table AGENDA")
para("Règles de normalisation. — Le modèle obtenu respecte les trois premières formes "
     "normales : chaque attribut est atomique (1FN), tout attribut non-clé dépend de "
     "l'intégralité de la clé (2FN) et il n'existe pas de dépendance transitive entre attributs "
     "non-clés (3FN). Cette normalisation garantit l'absence de redondance et préserve "
     "l'intégrité des données lors des mises à jour.")
para("Conclusion partielle. — La modélisation UML/Merise fournit une représentation rigoureuse "
     "du système. Le chapitre suivant en déduit l'architecture technique.")
br()

# ---- Chapitre 5
h1("CHAPITRE 5 : ARCHITECTURE DE LA SOLUTION")
h2("5.1. Architecture logique trois-tiers")
para("L'application repose sur une architecture trois-tiers (ou multi-couches), qui sépare "
     "nettement les responsabilités : la couche présentation (interface utilisateur), la "
     "couche métier (logique applicative et services) et la couche données (persistance). "
     "Cette séparation favorise la maintenabilité, la réutilisation et l'évolutivité.")
figure("architecture.png", "Architecture logique trois-tiers de l'application")
para("La couche présentation est assurée par une application monopage (SPA) développée avec "
     "Angular et la bibliothèque de composants PrimeNG. Elle communique avec le serveur via des "
     "requêtes HTTP au format JSON. La couche métier expose une API REST construite avec "
     "FastAPI, qui orchestre la logique applicative, applique les règles de gestion, gère "
     "l'authentification et délègue certains traitements lourds à des tâches asynchrones. La "
     "couche données s'appuie sur l'ORM SQLAlchemy et le SGBD relationnel PostgreSQL.")
h2("5.2. Style d'architecture REST")
para("L'interface entre le frontend et le backend respecte les principes de l'architecture "
     "REST : les ressources (dossiers, clients, paiements…) sont identifiées par des URL, et "
     "les opérations s'effectuent au moyen des verbes HTTP standard — GET pour la lecture, POST "
     "pour la création, PUT/PATCH pour la mise à jour et DELETE pour la suppression. Les "
     "échanges sont stateless : chaque requête transporte le jeton d'authentification "
     "nécessaire.")
h2("5.3. Architecture physique et déploiement")
para("Sur le plan physique, la solution se répartit entre le poste client (navigateur), le "
     "serveur d'application (serveur web Uvicorn exécutant FastAPI, worker Celery et broker de "
     "messages) et le serveur de base de données PostgreSQL. Le diagramme de déploiement "
     "ci-dessous illustre cette répartition.")
figure("deploiement.png", "Diagramme de déploiement de la solution")
h2("5.4. Sécurité de l'architecture")
para("La sécurité est traitée de manière transversale. L'authentification repose sur des "
     "jetons JWT signés, transmis à chaque requête. Les mots de passe sont hachés au moyen de "
     "l'algorithme bcrypt et ne sont jamais stockés en clair. Le contrôle d'accès basé sur les "
     "rôles (RBAC) restreint l'usage de chaque fonctionnalité selon le profil de l'utilisateur. "
     "Enfin, un journal d'audit enregistre les actions sensibles, garantissant la traçabilité.")
para("Conclusion partielle. — L'architecture définie offre un cadre robuste et évolutif. La "
     "troisième partie décrit sa réalisation concrète.")
br()

# ============================================================== PARTIE III
para("TROISIÈME PARTIE", bold=True, align="c", size=16, color=INDIGO)
para("RÉALISATION ET MISE EN ŒUVRE", bold=True, align="c", size=14, color=NAVY)
doc.add_paragraph()
para("Cette dernière partie présente les outils et modules informatiques mobilisés, décrit "
     "l'implémentation et les interfaces de l'application, puis aborde les tests, la sécurité "
     "et le déploiement.", italic=True, align="c")
br()

# ---- Chapitre 6
h1("CHAPITRE 6 : ENVIRONNEMENT ET MODULES INFORMATIQUES")
para("Ce chapitre détaille l'ensemble des langages, frameworks, bibliothèques et outils "
     "(modules informatiques) utilisés pour concevoir et réaliser l'application.")
h2("6.1. Langages de programmation")
para("Python a été retenu comme langage principal côté serveur pour sa lisibilité, sa "
     "richesse en bibliothèques et sa productivité. Côté client, TypeScript — sur-ensemble "
     "typé de JavaScript — a été utilisé pour bénéficier d'un typage statique fiable. Les "
     "langages HTML5 et CSS3 (via SCSS) assurent la structure et la mise en forme des pages.")
h2("6.2. Technologies du backend")
table(["Module / outil", "Rôle dans le projet"], [
    ["FastAPI 0.104", "Framework web Python pour exposer l'API REST, avec documentation OpenAPI automatique"],
    ["Uvicorn", "Serveur ASGI exécutant l'application FastAPI"],
    ["SQLAlchemy", "ORM assurant la correspondance objet-relationnel et l'accès aux données"],
    ["Pydantic 2.5", "Validation et sérialisation des données échangées (schémas)"],
    ["Alembic", "Gestion des migrations de la base de données"],
    ["python-jose", "Génération et vérification des jetons JWT"],
    ["passlib / bcrypt", "Hachage sécurisé des mots de passe"],
    ["Celery + RabbitMQ", "Exécution de tâches asynchrones (envois, traitements longs)"],
    ["python-multipart / aiofiles", "Gestion des téléversements et fichiers"],
    ["psycopg", "Pilote de connexion à PostgreSQL"],
], "Modules et bibliothèques du backend")
para("FastAPI mérite une attention particulière : ce framework moderne et performant repose "
     "sur les annotations de type de Python, génère automatiquement une documentation "
     "interactive (Swagger UI) et s'appuie sur Pydantic pour valider rigoureusement les données "
     "en entrée et en sortie. Cette approche réduit les erreurs et accélère le développement.")
h2("6.3. Technologies du frontend")
table(["Module / outil", "Rôle dans le projet"], [
    ["Angular", "Framework structurant l'application monopage (composants, routage, services)"],
    ["TypeScript", "Langage de développement typé du frontend"],
    ["PrimeNG", "Bibliothèque de composants d'interface (tables, formulaires, boîtes de dialogue)"],
    ["Chart.js (p-chart)", "Représentation graphique des statistiques du tableau de bord"],
    ["SCSS", "Feuilles de style structurées (thème professionnel global)"],
    ["RxJS", "Programmation réactive pour la gestion des flux de données HTTP"],
], "Modules et bibliothèques du frontend")
para("Angular est un framework structurant développé par Google, fondé sur les composants. "
     "Chaque écran de l'application correspond à un composant autonome regroupant son gabarit "
     "(template HTML), sa logique (classe TypeScript) et son style. Le routage assure la "
     "navigation entre les pages, tandis que les services centralisent les appels à l'API et le "
     "partage d'état. La bibliothèque RxJS, fondée sur la programmation réactive, gère "
     "élégamment l'asynchronisme des requêtes HTTP.")
para("PrimeNG enrichit Angular d'une collection de composants d'interface prêts à l'emploi et "
     "de qualité professionnelle : tableaux paginés et filtrables, formulaires, listes "
     "déroulantes, calendriers, boîtes de dialogue, notifications. Son usage a "
     "considérablement accéléré la réalisation d'interfaces homogènes. La bibliothèque Chart.js, "
     "intégrée via le composant graphique de PrimeNG, produit les diagrammes du tableau de bord "
     "(camemberts, histogrammes).")
h2("6.4. Système de gestion de base de données")
para("PostgreSQL a été choisi comme SGBD relationnel pour sa robustesse, sa conformité au "
     "standard SQL, son respect des propriétés ACID et sa gratuité. Il garantit l'intégrité "
     "référentielle des données grâce aux clés étrangères et aux contraintes définies dans le "
     "modèle physique.")
h2("6.5. Outils de développement et de gestion de projet")
table(["Outil", "Usage"], [
    ["Visual Studio Code", "Éditeur de code principal"],
    ["Git / GitHub", "Gestion de versions et collaboration"],
    ["Postman / Swagger UI", "Test et documentation des points d'accès de l'API"],
    ["pgAdmin / DBeaver", "Administration de la base PostgreSQL"],
    ["Node.js / npm", "Environnement d'exécution et gestion des paquets du frontend"],
], "Outils de développement utilisés")
para("Conclusion partielle. — Le socle technique étant présenté, le chapitre suivant décrit "
     "l'implémentation et les interfaces obtenues.")
br()

# ---- Chapitre 7
h1("CHAPITRE 7 : IMPLÉMENTATION ET PRÉSENTATION DE L'APPLICATION")
h2("7.1. Organisation du code")
para("Le projet est organisé en deux dépôts complémentaires : un backend et un frontend. Le "
     "backend suit une structure modulaire séparant les modèles (entités persistantes), les "
     "schémas (validation), les points d'accès (endpoints de l'API), les services (logique "
     "métier) et la configuration. Le frontend organise le code en composants, services et "
     "pages, un composant étant dédié à chaque module fonctionnel.")
h3("7.1.1. Exemple de définition d'un modèle (backend)")
code_block(
"""class Dossier(Base):
    __tablename__ = "dossiers"
    id = Column(Integer, primary_key=True, index=True)
    numero_dossier = Column(String(50), unique=True, nullable=False)
    objet = Column(String(255), nullable=False)
    type_dossier = Column(Enum(TypeDossier), nullable=False)
    statut = Column(Enum(StatutDossier), default=StatutDossier.NOUVEAU)
    date_creation = Column(DateTime, default=datetime.utcnow)
    id_client = Column(Integer, ForeignKey("clients.id"))
    id_user = Column(Integer, ForeignKey("utilisateurs.id"))""")
h3("7.1.2. Exemple de point d'accès de l'API (statistiques)")
code_block(
"""@router.get("/dashboard")
def get_dashboard_stats(db: Session = Depends(deps.get_db)):
    total_dossiers = db.query(func.count(Dossier.id)).scalar() or 0
    paiements_mois = db.query(func.sum(Paiement.montant))\\
        .filter(Paiement.date_paiement >= debut_mois).scalar() or 0
    return {"kpis": {"total_dossiers": total_dossiers,
                     "paiements_mois": float(paiements_mois)}}""")
h2("7.2. Présentation des interfaces")
para("Les interfaces ont été conçues selon une charte graphique professionnelle et cohérente, "
     "déclinant une palette de couleurs (bleu marine, indigo, vert, orange) et des composants "
     "uniformisés sur l'ensemble des pages.")
h3("7.2.1. Écran de connexion")
para("L'accès à l'application débute par un écran d'authentification sécurisé où "
     "l'utilisateur saisit son adresse électronique et son mot de passe.")
figure("mockup_login.png", "Interface — écran de connexion")
h3("7.2.2. Tableau de bord")
para("Une fois connecté, l'utilisateur accède au tableau de bord, véritable centre de "
     "pilotage du cabinet. Il présente des indicateurs clés (nombre de dossiers, encaissements "
     "du mois, nombre de clients, rendez-vous), des graphiques de répartition, une carte "
     "d'activité (heatmap) et un fil des dernières actions, le tout alimenté en temps réel par "
     "l'API et adapté au rôle de l'utilisateur connecté.")
figure("mockup_dashboard.png", "Interface — tableau de bord décisionnel")
para("Les éléments affichés s'adaptent dynamiquement aux droits de l'utilisateur : un "
     "secrétaire ne verra pas les indicateurs financiers réservés à l'huissier ou à "
     "l'administrateur. Le tableau de bord exploite des données réelles issues de la base : "
     "répartition des dossiers par type et par statut, comparatif du mois en cours par rapport "
     "au mois précédent, et activité quotidienne reconstituée à partir du journal d'audit.")
h2("7.3. Présentation détaillée des modules réalisés")
para("L'application met en œuvre l'ensemble des modules spécifiés lors de l'analyse. Chacun "
     "fait l'objet d'une page (ou d'un ensemble de pages) dédiée, accessible depuis le menu "
     "latéral en fonction des droits de l'utilisateur. Nous présentons ci-après chaque module, "
     "ses données, ses opérations et son ergonomie.")

h3("7.3.1. Module de gestion des dossiers")
para("Le module Dossiers est le pivot de l'application, car la quasi-totalité des autres "
     "objets (parties, actes, paiements, documents, rendez-vous) gravitent autour du dossier. "
     "L'interface se compose d'une barre d'outils (titre, recherche globale et boutons "
     "d'action), d'un tableau paginé et filtrable listant les dossiers, et d'une boîte de "
     "dialogue de saisie. Chaque dossier porte un numéro unique, un objet, un type (recouvrement, "
     "constat, signification, exécution), un statut (nouveau, en cours, terminé, archivé, "
     "annulé), une date de création, un client rattaché et l'utilisateur créateur.")
para("Les opérations disponibles couvrent la création, la consultation détaillée, la "
     "modification, le changement de statut, la suppression (réservée aux profils habilités) et "
     "l'export PDF de la liste. Un système de filtres permet de retrouver instantanément un "
     "dossier par son numéro, son objet, son statut ou son client, répondant directement à la "
     "lenteur de recherche identifiée dans l'existant. Toute opération d'écriture est "
     "enregistrée dans le journal d'audit.")

h3("7.3.2. Module de gestion des clients")
para("Le module Clients centralise la fiche de chaque donneur d'ordre du cabinet. Il distingue "
     "les personnes physiques (particuliers) des personnes morales (entreprises, banques, "
     "institutions), grâce à un champ « type de client ». La fiche comporte les coordonnées "
     "complètes : nom ou raison sociale, adresse, téléphone, courriel. Le module permet la "
     "création, la modification, la suppression et l'export, ainsi que la recherche rapide. Les "
     "données saisies une seule fois sont réutilisées dans tous les dossiers du client, "
     "supprimant la redondance constatée précédemment.")

h3("7.3.3. Module de gestion des parties")
para("Une partie désigne toute personne impliquée dans un dossier en une qualité déterminée : "
     "demandeur, défendeur, requérant, tiers saisi, etc. Le module Parties associe à chaque "
     "dossier la liste de ses parties avec leur rôle procédural et leurs coordonnées. Cette "
     "distinction entre « client » (qui mandate le cabinet) et « partie » (qui intervient dans "
     "l'affaire) reflète fidèlement la réalité du métier d'huissier.")

h3("7.3.4. Module de gestion des actes")
para("Le module Actes assure le suivi des actes juridiques rattachés à un dossier : type de "
     "l'acte, date, résultat ou issue, et observations. Il offre une traçabilité de l'activité "
     "procédurale et permet de constituer l'historique des diligences accomplies. Les actes "
     "sont accessibles aux clercs et aux huissiers, conformément à la matrice des droits.")

h3("7.3.5. Module de gestion des paiements")
para("Le module Paiements enregistre les sommes perçues par le cabinet : provisions, "
     "émoluments, frais et soldes. Chaque paiement précise son montant (en FCFA), son type, son "
     "mode de règlement (espèces, virement, chèque, mobile money) et sa date. Le module calcule "
     "automatiquement des totaux (total encaissé, encaissements du mois) affichés sous forme de "
     "cartes statistiques. Réservé aux profils financiers (huissier, administrateur), il "
     "alimente directement les indicateurs du tableau de bord et le comparatif mensuel.")

h3("7.3.6. Module de planification (agenda)")
para("Le module Agenda répond au besoin de planification partagée des activités. Il permet "
     "d'enregistrer des rendez-vous, audiences et opérations de signification, avec un titre, "
     "un type, une date de début et de fin, un lieu et un éventuel rattachement à un dossier. "
     "Une vue hebdomadaire offre une lecture rapide des activités à venir ; le tableau de bord "
     "affiche les rendez-vous du jour et de la semaine. Ce module remédie à l'absence d'agenda "
     "central relevée dans l'existant.")

h3("7.3.7. Modules documents et archives")
para("Le module Documents permet de joindre à chaque dossier les pièces numériques (actes "
     "scannés, justificatifs, courriers) en précisant leur nom, leur type et leur date d'ajout. "
     "Les fichiers sont stockés de manière organisée et téléchargeables à la demande. Le module "
     "Archives prend le relais pour les dossiers clôturés : il consigne la raison et la date "
     "d'archivage et permet de retrouver aisément un dossier archivé, garantissant la "
     "conservation et la sécurité des informations sur le long terme.")

h3("7.3.8. Modules de sécurité : utilisateurs, rôles et audit")
para("Le module Utilisateurs permet à l'administrateur de créer et gérer les comptes des "
     "collaborateurs. Le module Rôles présente les rôles applicatifs et les permissions "
     "associées, et permet d'attribuer un rôle à chaque utilisateur. Le module Audit affiche le "
     "journal horodaté de toutes les actions sensibles (création, modification, suppression, "
     "connexion), précisant l'auteur, l'action, l'entité concernée et la date. Ces trois modules "
     "constituent le socle de sécurité et de traçabilité de l'application.")

h3("7.3.9. Tableau de bord décisionnel")
para("Présenté plus haut, le tableau de bord agrège l'ensemble des informations pour offrir "
     "une vision synthétique de l'activité. Il est entièrement alimenté par des données réelles "
     "issues de la base : indicateurs clés, graphiques de répartition, carte d'activité "
     "reconstituée à partir du journal d'audit, fil des dernières actions et comparatif du mois "
     "courant par rapport au mois précédent. Son contenu s'adapte au rôle de l'utilisateur "
     "connecté.")
para("Conclusion partielle. — L'implémentation a permis de concrétiser l'ensemble des besoins "
     "exprimés. Chaque module répond directement à une faiblesse identifiée dans l'existant. Le "
     "dernier chapitre traite des tests, de la sécurité et du déploiement.")
br()

# ---- Chapitre 8
h1("CHAPITRE 8 : TESTS, SÉCURITÉ ET DÉPLOIEMENT")
h2("8.1. Stratégie de tests")
para("La qualité du logiciel a été contrôlée à plusieurs niveaux. Des tests unitaires ont "
     "vérifié le comportement des fonctions critiques (validations, calculs). Des tests "
     "d'intégration ont contrôlé le bon fonctionnement des points d'accès de l'API à l'aide de "
     "Postman et de la documentation Swagger générée automatiquement. Enfin, des tests "
     "fonctionnels manuels ont validé les scénarios d'utilisation de bout en bout, du point de "
     "vue des différents profils.")
table(["Type de test", "Objectif", "Outil"], [
    ["Unitaire", "Vérifier une fonction isolée", "pytest"],
    ["Intégration", "Vérifier les endpoints de l'API", "Postman / Swagger"],
    ["Fonctionnel", "Valider les scénarios métier", "Tests manuels"],
    ["Compilation", "Garantir l'absence d'erreurs", "ng build / py_compile"],
], "Stratégie de tests adoptée")
h2("8.2. Sécurité applicative")
para("La sécurité constitue une exigence forte compte tenu de la confidentialité des données "
     "juridiques. Les mesures mises en place sont les suivantes :")
for s in ["authentification par jeton JWT signé, à durée de validité limitée ;",
          "hachage des mots de passe par l'algorithme bcrypt ;",
          "contrôle d'accès basé sur les rôles (RBAC) appliqué côté serveur et côté client ;",
          "validation systématique des données entrantes via Pydantic ;",
          "journalisation des actions sensibles dans un journal d'audit horodaté ;",
          "communication chiffrée par HTTPS en production."]:
    bullet(s)
h2("8.3. Déploiement")
para("Le déploiement de la solution suit une procédure reproductible : configuration des "
     "variables d'environnement (chaîne de connexion à la base, clé secrète JWT), application "
     "des migrations Alembic pour créer le schéma, lancement du serveur Uvicorn pour l'API et "
     "du worker Celery, puis construction (build) du frontend Angular et publication des "
     "fichiers statiques. L'application est ainsi accessible aux utilisateurs depuis leur "
     "navigateur.")
para("Conclusion partielle. — Les tests et les dispositifs de sécurité confirment la fiabilité "
     "de la solution livrée, prête à être déployée au sein du cabinet.")
br()

# ============================================================== CONCLUSION
h1("CONCLUSION GÉNÉRALE ET PERSPECTIVES")
para("Au terme de ce travail, nous avons conçu et implémenté une application web complète de "
     "gestion des dossiers et de planification des activités pour un cabinet d'huissier de "
     "justice. Partant d'un constat de gestion essentiellement manuelle et dispersée, source "
     "de lenteurs et de risques d'erreurs, nous avons mené une analyse rigoureuse des besoins, "
     "puis une modélisation appuyée sur UML et Merise, avant de réaliser la solution selon une "
     "architecture trois-tiers moderne associant Angular, FastAPI et PostgreSQL.")
para("L'application livrée centralise l'information, automatise le suivi des dossiers, des "
     "actes et des paiements, offre une planification partagée des activités, sécurise l'accès "
     "par une gestion fine des rôles et assure la traçabilité grâce à un journal d'audit. Un "
     "tableau de bord décisionnel, alimenté par des données réelles, fournit au cabinet une "
     "vision consolidée et en temps réel de son activité.")
para("Ce stage a constitué une expérience particulièrement enrichissante, tant sur le plan "
     "technique que professionnel. Il nous a permis de mettre en pratique les connaissances "
     "acquises durant notre formation, de découvrir des technologies modernes et de mieux "
     "appréhender les réalités d'un environnement professionnel et les exigences d'un métier "
     "juridique.")
para("Perspectives. — Plusieurs améliorations peuvent être envisagées pour enrichir la "
     "solution :")
for p_ in ["la mise en place d'un système d'alertes automatiques sur les échéances et délais légaux ;",
           "l'ajout d'un module de génération automatique des actes à partir de modèles ;",
           "le développement d'une application mobile pour les opérations sur le terrain ;",
           "l'intégration d'une signature électronique des actes ;",
           "l'enrichissement du module statistique par des analyses prédictives ;",
           "la sauvegarde automatisée et la haute disponibilité de la base de données."]:
    bullet(p_)
br()

# ============================================================== BIBLIOGRAPHIE
h1("BIBLIOGRAPHIE ET WEBOGRAPHIE")
para("Ouvrages et articles", bold=True, color=NAVY)
for ref in [
    "TARDIEU H., ROCHFELD A., COLLETTI R., La méthode Merise : Principes et outils, Éditions d'Organisation, Paris.",
    "ROQUES P., VALLÉE F., UML 2 en action : de l'analyse des besoins à la conception, Eyrolles, Paris.",
    "FOWLER M., UML Distilled: A Brief Guide to the Standard Object Modeling Language, Addison-Wesley.",
    "GAUDEL M.-C., MARRE B., SCHLIENGER F., Précis de génie logiciel, Masson.",
    "DI GALLO F., Méthodologie des systèmes d'information — Merise, support de cours.",
]:
    p = doc.add_paragraph(ref, style="List Number")
para("Webographie", bold=True, color=NAVY)
for ref in [
    "Documentation officielle FastAPI — https://fastapi.tiangolo.com (consulté en 2026).",
    "Documentation officielle Angular — https://angular.dev (consulté en 2026).",
    "Documentation officielle PostgreSQL — https://www.postgresql.org/docs (consulté en 2026).",
    "Documentation PrimeNG — https://primeng.org (consulté en 2026).",
    "Documentation SQLAlchemy — https://docs.sqlalchemy.org (consulté en 2026).",
    "Spécification UML, Object Management Group — https://www.omg.org/spec/UML (consulté en 2026).",
]:
    p = doc.add_paragraph(ref, style="List Number")
br()

# ============================================================== ANNEXES
h1("ANNEXES")
h2("Annexe A : Liste complète des entités de la base de données")
table(["Table", "Rôle"], [
    ["cabinets", "Informations sur le cabinet"],
    ["utilisateurs", "Comptes et rôles des collaborateurs"],
    ["dossiers", "Dossiers / affaires traitées"],
    ["clients", "Clients du cabinet"],
    ["parties", "Parties impliquées dans les dossiers"],
    ["actes", "Actes juridiques"],
    ["paiements", "Paiements et provisions"],
    ["agenda", "Rendez-vous et activités planifiées"],
    ["documents", "Pièces jointes aux dossiers"],
    ["archives", "Dossiers archivés"],
    ["affectation_dossier", "Affectation des dossiers aux collaborateurs"],
    ["audit_logs", "Journal d'audit des actions"],
], "Liste des tables de la base de données")
h2("Annexe B : Extrait de configuration de la base de données")
code_block(
"""DATABASE_URL = "postgresql://postgres:postgres@localhost:5432/dossier_db"
engine = create_engine(DATABASE_URL, pool_pre_ping=True)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)""")
h2("Annexe C : Rôles applicatifs et permissions (RBAC)")
table(["Rôle", "Accès principaux"], [
    ["ADMIN", "Tous les modules, gestion des utilisateurs et audit"],
    ["HUISSIER", "Dossiers, actes, paiements, affectations, archives, tableau de bord"],
    ["CLERC", "Dossiers, actes, parties, clients"],
    ["SECRETAIRE", "Clients, agenda, documents"],
    ["ASSISTANT", "Dossiers (consultation et suivi), agenda"],
], "Matrice des rôles et permissions")
br()

# ============================================================== TABLE DES MATIÈRES
toc_field("TABLE DES MATIÈRES")

OUT = os.path.join(HERE, "Rapport_de_stage_cabinet_huissier.docx")
doc.save(OUT)
print("Document généré :", OUT)
print("Figures:", FIG_N[0], "| Tableaux:", TAB_N[0])
